from pathlib import Path


def list_files_in_folder(folder_path: str | Path) -> list[str]:
    """
    Return a list of all file names (not subdirectories) in the given folder.
    Returns an empty list if the path does not exist or is not a directory.
    """
    path = Path(folder_path)
    if not path.is_dir():
        return []
    return sorted(
        p.name for p in path.iterdir() if p.is_file()
    )


def _read_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n\n".join(parts)


def read_file_as_text(file_path: str | Path) -> str:
    """
    Return textual content for .txt, .md, .pdf, and .docx; raises on unsupported type.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md"):
        return _read_plain_text(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(
        f"Unsupported file type {suffix!r}; use .txt, .md, .pdf, or .docx"
    )
